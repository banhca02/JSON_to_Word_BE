from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_table_border(table):
    tbl = table._tbl  # Lấy đối tượng XML gốc của bảng
    tblPr = tbl.tblPr

    borders = OxmlElement('w:tblBorders')

    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')      # Kiểu viền: single, double, dashed, dotted...
        border.set(qn('w:sz'), '6')           # Độ dày (1/8 pt), 12 = 1.5pt
        border.set(qn('w:space'), '1.5')         # Khoảng cách
        border.set(qn('w:color'), '000000')    # Màu viền (hex)

        borders.append(border)

    tblPr.append(borders)

    # Thêm khoảng cách giữa các ô
    cell_spacing = OxmlElement('w:tblCellSpacing')
    cell_spacing.set(qn('w:w'), '15')         
    cell_spacing.set(qn('w:type'), 'dxa')
    tblPr.append(cell_spacing)

def set_column_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def center_align_column(table, col_idx):
    for row in table.rows:
        cell = row.cells[col_idx]
        # Căn ngang nội dung (paragraph)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Căn dọc ô
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def center_align_row(table, row_idx):
    for col in table.columns:
        cell = col.cells[row_idx]
        # Căn ngang nội dung (paragraph)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Căn dọc ô
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def create_file_word(json_data, filepath):
    # Tạo file Word
    doc = Document()

    title_style = doc.styles['Title']
    title_font = title_style.font
    title_font.name = 'Times New Roman'
    title_font.size = Pt(26)
    title_font.bold = True

    heading1_style = doc.styles['Heading 1']
    heading1_font = heading1_style.font
    heading1_font.name = 'Times New Roman'
    heading1_font.size = Pt(14)
    heading1_font.bold = True

    heading2_style = doc.styles['Heading 2']
    heading2_font = heading2_style.font
    heading2_font.name = 'Times New Roman'
    heading2_font.size = Pt(13)
    heading2_font.bold = True

    # Đổi font mặc định cho style 'Normal'
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Nếu cần đổi thêm các style khác (đặc biệt cho bảng):
    table_style = doc.styles['Table Grid']
    table_font = table_style.font
    table_font.name = 'Times New Roman'
    table_font.size = Pt(12)

    # Bảng tóm tắt các lỗ hổng
    doc.add_heading("Summary of Findings", 0)

    summary_table = doc.add_table(rows=1, cols=4)
    summary_table.style = 'Table Grid'

    summary_table.autofit = False 
    hdr_cells = summary_table.rows[0].cells

    headers = ['Code', 'Title', 'Risk Rating', 'Affected Function(s)']
    for i, text in enumerate(headers):
        paragraph = hdr_cells[i].paragraphs[0]
        run = paragraph.add_run(text)
        run.bold = True

    # Lặp qua dữ liệu để điền bảng
    for item in json_data:
        row_cells = summary_table.add_row().cells
        row_cells[0].text = item.get('code', '')
        row_cells[1].text = item.get('name', '')

        # Chỉnh màu cho rating
        rating_cell = row_cells[2].paragraphs[0]
        rating = item.get('rating', '')
        rating_run = rating_cell.add_run(rating)
        rating_run.bold = True
        rating_lower = rating.lower()
        if rating_lower == 'low':
            rating_run.font.color.rgb = RGBColor(0, 176, 80)      # Xanh lá
        elif rating_lower == 'medium':
            rating_run.font.color.rgb = RGBColor(255, 192, 0)     # Cam
        elif rating_lower == 'high':
            rating_run.font.color.rgb = RGBColor(255, 0, 0)       # Đỏ
        elif rating_lower == 'critical':
            rating_run.font.color.rgb = RGBColor(128, 0, 128)     # Tím

        affected = item.get('affected_items', [])
        functions = [a.get('affected_function', '') for a in affected if a.get('affected_function')]
        row_cells[3].text = ', '.join(functions)

    # Chỉnh độ rộng cột
    set_column_width(summary_table, 0, 2.0)  # Cột Code
    set_column_width(summary_table, 1, 6.0)  # Cột Title
    set_column_width(summary_table, 2, 2.5)  # Cột Rating
    set_column_width(summary_table, 3, 6.0)  # Cột Affected Function(s)

    # Căn giữa cột
    center_align_column(summary_table, 0)  # Cột Code
    center_align_column(summary_table, 2)  # Cột Rating

    # Căn giữa hàng
    center_align_row(summary_table, 0)  # Hàng tiêu đề

    set_table_border(summary_table)


    ###################################################################
    doc.add_paragraph("") 

    # Báo cáo chi tiết
    doc.add_heading('Vulnerability Details', 0)

    for item in json_data:
        doc.add_heading(item.get("name", "No Name"), level=1)
        doc.add_paragraph(f"Code: {item.get('code', '')}")
        doc.add_paragraph(f"Rating: {item.get('rating', '')}")

        doc.add_paragraph('Description:', style='Heading 2')
        doc.add_paragraph(item.get('description', ''))

        doc.add_paragraph('Impact:', style='Heading 2')
        doc.add_paragraph(item.get('impact', ''))

        doc.add_paragraph('Recommendation:', style='Heading 2')
        doc.add_paragraph(item.get('recommendation', ''))

        doc.add_paragraph('Affected Items:', style='Heading 2')
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Function'
        hdr_cells[1].text = 'URI'

        for affected in item.get('affected_items', []):
            row_cells = table.add_row().cells
            row_cells[0].text = affected.get('affected_function', '')
            row_cells[1].text = affected.get('affected_uri', '')

        doc.add_paragraph("") 

    doc.save(filepath)

